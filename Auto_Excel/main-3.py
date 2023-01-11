from docx import Document
from openpyxl import load_workbook
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

load_wb = load_workbook('수료증명단.xlsx');
load_ws = load_wb.active;

name_list = [];
birthday_list = [];
ho_list = [];

for i in range(1, load_ws.max_row + 1):
    name_list.append(load_ws.cell(i, 1).value);
    birthday_list.append(load_ws.cell(i, 2).value);
    ho_list.append(load_ws.cell(i, 3).value);

print(name_list);
print(birthday_list);
print(ho_list);


for i in range(len(name_list)):
    doc = docx.Document('수료증양식.docx');
    style = doc.styles["Normal"];
    style.font.name = "나눔고딕";
    style._element.rPr.rFonts.set(qn('w:eastAsia'), "나눔고딕");
    style.font.size = docx.shared.Pt(12);

    para = doc.add_paragraph();  # 문단생성
    run = para.add_run('\n\n\n');
    run = para.add_run('            제' + ho_list[i] +'  호\n');
    run.font.name = "나눔고딕";
    style._element.rPr.rFonts.set(qn('w:eastAsia'), "나눔고딕");
    run.font.size = docx.shared.Pt(20);

    para = doc.add_paragraph();
    run = para.add_run('\n\n');
    run = para.add_run('수료증\n');
    run.font.name = "나눔고딕";
    run.bold = True;
    style._element.rPr.rFonts.set(qn('w:eastAsia'), "나눔고딕");
    run.font.size = docx.shared.Pt(40);
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER;  # 중앙 정렬

    para = doc.add_paragraph();
    run = para.add_run('\n');
    run = para.add_run('        성 명 : ' + name_list[i] + '\n');
    run.font.name = "나눔고딕";
    style._element.rPr.rFonts.set(qn('w:eastAsia'), "나눔고딕");
    run.font.size = docx.shared.Pt(20);
    run = para.add_run('        생 년 월 일 : ' + birthday_list[i] +'\n');
    run.font.name = '나눔고딕';
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "나눔고딕");
    run.font.size = docx.shared.Pt(20);
    run = para.add_run("        교 육 과 정: 파이썬과 40개의 작품들\n");
    run.font.name = "나눔고딕";
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "나눔고딕");
    run.font.size = docx.shared.Pt(20);
    run = para.add_run("        교 육 날 짜: 2021.08.05~2021.09.09\n");
    run.font.name = "나눔고딕";
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "나눔고딕");
    run.font.size = docx.shared.Pt(20);

    para = doc.add_paragraph();
    run = para.add_run('\n\n');
    run = para.add_run('        위 사람은 파이썬과 40개의 자품들 교육과정을\n');
    run.font.name = "나눔고딕";
    style._element.rPr.rFonts.set(qn('w:eastAsia'), "나눔고딕");
    run.font.size = docx.shared.Pt(20);
    run = para.add_run("         이수하였으므로 이 증서를 수여 합니다.\n");
    run.font.name = "나눔고딕";
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "나눔고딕");
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph();
    run = para.add_run('\n\n\n');
    run = para.add_run('2021.09.19\n');
    run.font.name = "나눔고딕";
    style._element.rPr.rFonts.set(qn('w:eastAsia'), "나눔고딕");
    run.font.size = docx.shared.Pt(20);
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER;  # 중앙 정렬

    para = doc.add_paragraph();
    run = para.add_run('\n\n\n');
    run = para.add_run('파이썬교육기관장\n');
    run.font.name = "나눔고딕";
    run.bold = True;
    style._element.rPr.rFonts.set(qn('w:eastAsia'), "나눔고딕");
    run.font.size = docx.shared.Pt(20);
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER;  # 중앙 정렬

    doc.save(name_list[i] + '.docx');
    convert(name_list[i] + '.docx', name_list[i] + '.pdf');